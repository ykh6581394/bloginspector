# -*- coding: utf-8 -*-
"""
Created on Mon Jun 17 07:38:58 2024

@author: 유국현
"""

import streamlit as st
import pandas as pd
import io
import os
import zipfile
from docx import Document
import time
import docx2txt
import shutil
#import aspose.words as aw
#from sentence_transformers import SentenceTransformer, util

def KeyInspector(paras, keyword, sep):
    ek = keyword
    ek_result = []
    in_word = []
    no_word = []
    for e in range(len(ek)):
        if ek[e] in paras:
            ek_result.append("Pass")
            in_word.append(ek[e])
        else:
            ek_result.append("NoKey")
            no_word.append(ek[e])
    testDict = dict(zip(ek, ek_result))

    result = str(len(in_word)) + " in " +str(len(ek_result))
  
    for key, value in testDict.items():
        print('{} : {}'.format(key,value))
    result_dt = pd.DataFrame({sep+"_Keyword":ek,sep+"_KeyIn":ek_result})
    return in_word, no_word, result_dt, result




@st.cache_data
def convert_df(df):
   return df.to_csv(index=False).encode('utf-8-sig')


st.title("PRONE Blog Inspector")


st.sidebar.title('사용 메뉴얼')
select_ft = st.sidebar.selectbox(
    'Select Your Project',
    ['P&G / Braun','Pampers','C']
)
if select_ft == "P&G / Braun":
    st.sidebar.subheader("Keyword 검수 프로세스")
    st.sidebar.write("1. Keyword가 포함된 csv 파일 업로드")
    st.sidebar.write("2. 블로거 작성 Word 파일 업로드")
    st.sidebar.write("3. Keyword Extract 버튼 클릭")
    st.sidebar.write("4. [Download] 버튼 클릭")
    st.sidebar.write(" ")
    st.sidebar.subheader("Image 검수 프로세스")
    st.sidebar.write("1. 블로거 작성 Word 파일 업로드")
    st.sidebar.write("2. Image Extract 버튼 클릭")
    st.sidebar.write("3. [Download] 버튼 클릭")
    
elif select_ft == "Pampers":
    st.sidebar.subheader("Keyword 검수 프로세스")
    st.sidebar.write("1. 필수 키워드가 포함되어야 하는 라인 수")
    st.sidebar.write("2. Keyword가 포함된 csv 파일 업로드")
    st.sidebar.write("3. 블로거 작성 Word 파일 업로드")   
    st.sidebar.write("3. Keyword Extract 버튼 클릭")
    st.sidebar.write("4. [Download] 버튼 클릭")

    
elif select_ft == "C":
    st.sidebar.write("메뉴얼C")


tab1, tab2, tab3 = st.tabs(["P&G / Braun","Pampers", "Reply"])

with tab1:

    uploaded_files =  st.file_uploader("Upload your keyword file",type=['csv'],accept_multiple_files=False)
    docx_files     =  st.file_uploader("Upload your word file",type=['docx'],accept_multiple_files=True)
    
    if st.button("Keyword Extract"):
        key_all = pd.read_csv(uploaded_files)
        
        essential_keyword  = list(key_all["essential_Keyword"].dropna())
        selective_keyword  = list(key_all["selective_Keyword"].dropna())
        essential_tag      = list(key_all["tag_Keyword"].dropna())
        prohibited_keyword = list(key_all["prohibited_Keyword"].dropna())
    
        
        buf = io.BytesIO()
        with zipfile.ZipFile(buf,"x") as csv_zip:
            
            org_name       = []
            title_name     = []
            essential_res  = []
            essential_in   = []
            essential_no   = []
            selective_res  = []
            selective_in   = []
            selective_no   = []
            tag_res        = []
            tag_in   = []
            tag_no   = []
            prohibited_res = [] 
            prohibited_in   = []
            prohibited_no   = []

            for idx in docx_files:
                
                doc = Document(idx)
                file_name = idx.name
                org_name.append(file_name.split(".")[0])
                paras = []
                                
                nouse = ['','\ufeff']
                for para in doc.paragraphs:
                    if not para.text in nouse:
                        para = para.text.replace("\ufeff","")
                        paras.append(para)
                title = paras[0]
                title_name.append(title)
                paras = ''.join(paras)
                #st.write(paras)
               
                e_in, e_no, e_dt, e_res = KeyInspector(paras, essential_keyword, "essential")
                s_in, s_no, s_dt, s_res = KeyInspector(paras, selective_keyword, "selective")
                t_in, t_no, t_dt, t_res = KeyInspector(paras, essential_tag, "tag")
                p_in, p_no, p_dt, p_res = KeyInspector(paras, prohibited_keyword, "prohibited")
                essential_res.append(e_res)
                essential_in.append(e_in)
                essential_no.append(e_no)
                selective_res.append(s_res)
                selective_in.append(s_in)
                selective_no.append(s_no)
                tag_res.append(t_res)
                tag_in.append(t_in)
                tag_no.append(t_no)
                prohibited_res.append(p_res)
                prohibited_in.append(p_in)
                prohibited_no.append(p_no)
                
                result_dt = pd.concat([e_dt,s_dt,t_dt, p_dt], axis=1)
            
                csv_zip.writestr(title+"_"+file_name.split(".")[0]+".csv",result_dt.to_csv(index=False).encode('utf-8-sig'))
                
            summary_dt = pd.DataFrame({"filename" : org_name,
                                       "title" : title_name,
                                       "essential_result" : essential_res,
                                       "essential_in" : essential_in,
                                       "essential_no" : essential_no,
                                       "selective_result" : selective_res,
                                       "selective_in" : selective_in,
                                       "selective_no" : selective_no,
                                       "tag_result" : tag_res,
                                       "tag_in" : tag_in,
                                       "tag_no" : tag_no,
                                       "prohibited_result" : prohibited_res,
                                       "prohibited_in" : prohibited_in,
                                       "prohibited_no" : prohibited_no
                                       })
                
            csv_zip.writestr("summary.csv",summary_dt.to_csv(index=False).encode('utf-8-sig'))
                
                
            st.download_button(
                    label = "[Download] Download keyword zip",
                    data = buf.getvalue(),
                    file_name = "keyword_download.zip"
                    )
     
    if st.button("Image Extract"):
        
        img_path = "./result_image/"
        if os.path.exists(img_path):
            shutil.rmtree(img_path)

        os.mkdir(img_path)

        for idx in docx_files:
        
            doc = Document(idx)
            file_name = idx.name
            os.mkdir(img_path+file_name.split(".")[0])
            
            docx2txt.process(idx, img_path+file_name.split(".")[0])

        
        

        buf = io.BytesIO()
        with zipfile.ZipFile(buf,"x") as img_zip:
            for (path, dir, files) in os.walk(img_path):
                for file in files:
                    #if file.endswith('.jpeg') or file.endswith('.png'):
                    img_zip.write(os.path.join(path, file), compress_type=zipfile.ZIP_DEFLATED)

            st.write("Complete Image Extraction")
            
            st.download_button(
                label = "[Download] Download image zip",
                data = buf.getvalue(),
                file_name = "image.zip",
                mime = 'application/zip'
        
            )


with tab2:  
    start_line = st.text_input("필수 키워드가 포함되어야 하는 Line 수")
    uploaded_files2 =  st.file_uploader("Upload your keyword file2",type=['csv'],accept_multiple_files=False)
    docx_files2     =  st.file_uploader("Upload your word file2",type=['docx'],accept_multiple_files=True)
    
    if st.button("Keyword Extract2"):
        key_all = pd.read_csv(uploaded_files2)
        
        essential_keyword  = list(key_all["essential_Keyword"].dropna())
        selective_keyword  = list(key_all["selective_Keyword"].dropna())
        essential_tag      = list(key_all["tag_Keyword"].dropna())
        prohibited_keyword = list(key_all["prohibited_Keyword"].dropna())
        essential5_keyword = list(key_all["essential5_Keyword"].dropna())
        disc_keyword       = list(key_all["dis_Keyword"].dropna())
        
        buf = io.BytesIO()
        with zipfile.ZipFile(buf,"x") as csv_zip:
            
            #embedder = SentenceTransformer("jhgan/ko-sroberta-multitask")
            
            org_name       = []
            title_name     = []
            essential_res  = []
            essential_in   = []
            essential_no   = []
            selective_res  = []
            selective_in   = []
            selective_no   = []
            tag_res  = []
            tag_in   = []
            tag_no   = []
            prohibited_res  = [] 
            prohibited_in   = []
            prohibited_no   = []
            essential5_res  = []
            essential5_in   = []
            essential5_no   = []
            dis_res = []
            dis_in  = []
            dis_no  = []


            for idx in docx_files2:
                
                doc = Document(idx)
                file_name = idx.name
                org_name.append(file_name.split(".")[0])
                paras = []
                dis_paras = []                
                
                nouse = ['','\ufeff']
                for para in doc.paragraphs:
                    if not para.text in nouse:
                        para = para.text.replace("\ufeff","")
                        paras.append(para)
                    #print(para.text)
                    
                dis_dis = []
                dis_suc = []
                dis_fail = []
                for pp in paras:
                    for d in disc_keyword:
                        if d in pp:
                            dis_paras.append(pp)
                            if "*" in pp:
                                dis_dis.append("Pass")
                                dis_suc.append(pp)
                            else:
                                dis_dis.append("NoDis")
                                dis_fail.append(pp)

                dis_res.append(str(len(dis_suc)) + " in " +str(len(dis_dis)))
                dis_in.append(dis_suc)
                dis_no.append(dis_fail)
                
                title = paras[0]
                title_name.append(title)
                sel_paras = paras[:int(start_line)]
                paras = ''.join(paras)
                sel_paras = ''.join(sel_paras)
                #st.write(paras)
               
                e_in, e_no, e_dt, e_res = KeyInspector(paras, essential_keyword, "essential")
                s_in, s_no, s_dt, s_res = KeyInspector(paras, selective_keyword, "selective")
                t_in, t_no, t_dt, t_res = KeyInspector(paras, essential_tag, "tag")
                p_in, p_no, p_dt, p_res = KeyInspector(paras, prohibited_keyword, "prohibited")
                e5_in, e5_no, e5_dt, e5_res = KeyInspector(paras, prohibited_keyword, "essential5")
                
                essential_res.append(e_res)
                essential_in.append(e_in)
                essential_no.append(e_no)
                selective_res.append(s_res)
                selective_in.append(s_in)
                selective_no.append(s_no)
                tag_res.append(t_res)
                tag_in.append(t_in)
                tag_no.append(t_no)
                prohibited_res.append(p_res)
                prohibited_in.append(p_in)
                prohibited_no.append(p_no)
                essential5_res.append(e5_res)
                essential5_in.append(e5_in)
                essential5_no.append(e5_no)
                
                result_dt = pd.concat([e_dt,s_dt,t_dt, p_dt, e5_dt], axis=1)
            
                csv_zip.writestr(title+"_"+file_name.split(".")[0]+".csv",result_dt.to_csv(index=False).encode('utf-8-sig'))
            
            #embedder = SentenceTransformer("jhgan/ko-sroberta-multitask")
            #corpus_embeddings = embedder.encode(title_name, convert_to_tensor=True)
            
            
            #cos_scores = util.pytorch_cos_sim(corpus_embeddings, corpus_embeddings)
            #cos_scores = cos_scores.cpu()
            #title_sim = pd.DataFrame(cos_scores, columns = org_name, index = org_name)


            summary_dt = pd.DataFrame({"filename" : org_name,
                                       "title" : title_name,
                                       "essential_result" : essential_res,
                                       "essential_in" : essential_in,
                                       "essential_no" : essential_no,
                                       "selective_result" : selective_res,
                                       "selective_in" : selective_in,
                                       "selective_no" : selective_no,
                                       "tag_result" : tag_res,
                                       "tag_in" : tag_in,
                                       "tag_no" : tag_no,
                                       "prohibited_result" : prohibited_res,
                                       "prohibited_in" : prohibited_in,
                                       "prohibited_no" : prohibited_no,
                                       "essential5_result" : essential5_res,
                                       "essential5_in" : essential5_in,
                                       "essential5_no" : essential5_no,
                                       "dis_result" : dis_res,
                                       "dis_in" : dis_in,
                                       "dis_no" : dis_no,
                                       })
                
            csv_zip.writestr("summary.csv",summary_dt.to_csv(index=False).encode('utf-8-sig'))
            #csv_zip.writestr("title_similarity.csv",title_sim.to_csv().encode('utf-8-sig'))
                
            st.download_button(
                    label = "[Download] Download keyword zip",
                    data = buf.getvalue(),
                    file_name = "keyword_download.zip"
                    )
     
    if st.button("Image Extract2"):
        
        img_path = "./result_image/"
        if os.path.exists(img_path):
            shutil.rmtree(img_path)

        os.mkdir(img_path)
        
        
        for idx in docx_files:
        
            doc = Document(idx)
            file_name = idx.name
            os.mkdir(img_path+file_name.split(".")[0])
            
            docx2txt.process(idx, img_path+file_name.split(".")[0])


        buf = io.BytesIO()
        with zipfile.ZipFile(buf,"x") as img_zip:
            for (path, dir, files) in os.walk(img_path):
                for file in files:
                    #if file.endswith('.jpeg') or file.endswith('.png'):
                    img_zip.write(os.path.join(path, file), compress_type=zipfile.ZIP_DEFLATED)

            st.write("Complete Image Extraction")
            
            st.download_button(
                label = "[Download] Download image zip",
                data = buf.getvalue(),
                file_name = "image.zip",
                mime = 'application/zip'
        
            )
            

with tab3:

    reply_files   =  st.file_uploader("Upload your Reply",type=['xlsx'],accept_multiple_files=False)
    keyword_files =  st.file_uploader("Upload your keyword",type=['csv'],accept_multiple_files=False)
    
    if st.button("Reply Counter"):
    
        xl = pd.ExcelFile(reply_files)
        xl_sheet = xl.sheet_names
    
        key = list(pd.read_csv(keyword_files)["keyword"])
        
        influencer = []
        counter    = []
    
        for xx in range(len(xl_sheet)):
            reply_all = pd.read_excel(reply_files, sheet_name = xl_sheet[xx])
            reply_list = list(reply_all["Reply"])
            influencer.append(xl_sheet[xx])
            rep_each_counter = []
            for rep in reply_list:
                unit = 0
                for k in key:
                    if k in rep:
                        unit += 1
                rep_each_counter.append(unit)
            counter.append(sum(rep_each_counter))
        
        reply_dt = pd.DataFrame({"influencer" : influencer, "reply counter" : counter})
    
        csv = convert_df(reply_dt)
        st.download_button(
            "[Download] Press Download Reply",
            csv,
            "replyCounter.csv",
            key="download_csv")
        
        
        
        
        
        
        
        
        